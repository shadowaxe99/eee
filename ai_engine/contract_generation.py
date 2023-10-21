```python
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class ContractGenerator:
    def __init__(self, influencer, brand, terms):
        self.influencer = influencer
        self.brand = brand
        self.terms = terms
        self.document = Document()

    def generate_contract(self):
        self._add_title()
        self._add_parties()
        self._add_terms()
        self._add_signatures()
        return self._save_contract()

    def _add_title(self):
        title = self.document.add_paragraph()
        run = title.add_run('Partnership Agreement')
        run.bold = True
        run.font.size = Pt(24)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def _add_parties(self):
        self.document.add_paragraph(f"This agreement is made between {self.influencer} ('Influencer') and {self.brand} ('Brand').")

    def _add_terms(self):
        self.document.add_paragraph("The terms of the agreement are as follows:")
        for term in self.terms:
            self.document.add_paragraph(f"- {term}")

    def _add_signatures(self):
        self.document.add_paragraph("\n\n\n______________________     ______________________")
        self.document.add_paragraph(f"{self.influencer}     {self.brand}")
        self.document.add_paragraph("Signature     Signature")

    def _save_contract(self):
        filename = f"{self.influencer}_{self.brand}_contract.docx"
        self.document.save(filename)
        return filename

def generate_contract(influencer, brand, terms):
    generator = ContractGenerator(influencer, brand, terms)
    return generator.generate_contract()
```